import os
from typing import Dict, Any
import docx
from docx.shared import Pt, Inches


class StandardDOCXRenderer:
    """
    Standard DOCX Renderer generating clean, ATS-scannable Microsoft Word documents.
    """

    @staticmethod
    def _to_str(val: Any) -> str:
        if not val:
            return ""
        if isinstance(val, list):
            return ", ".join([str(x) for x in val])
        return str(val)

    @classmethod
    def render_docx(cls, doc_data: Dict[str, Any], output_path: str) -> str:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        doc = docx.Document()

        # Page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # Header
        header = doc_data.get("header", {})
        h_p = doc.add_paragraph()
        run = h_p.add_run(header.get("full_name", ""))
        run.bold = True
        run.font.size = Pt(18)

        c_p = doc.add_paragraph()
        c_p.add_run(f"Email: {header.get('email', '')} | Phone: {header.get('phone', 'N/A')} | Location: {header.get('location', 'N/A')}")

        # Summary
        doc.add_heading("PROFESSIONAL SUMMARY", level=2)
        doc.add_paragraph(doc_data.get("summary", ""))

        # Skills
        doc.add_heading("TECHNICAL SKILLS", level=2)
        skills_list = [s.get("name", "") if isinstance(s, dict) else str(s) for s in doc_data.get("skills", [])]
        doc.add_paragraph(", ".join(skills_list))

        # Projects
        projects = doc_data.get("projects", [])
        if projects:
            doc.add_heading("PROJECTS", level=2)
            for p in projects:
                p_p = doc.add_paragraph()
                r_title = p_p.add_run(p.get("name", ""))
                r_title.bold = True
                techs = cls._to_str(p.get("technologies"))
                desc = cls._to_str(p.get("description"))
                p_p.add_run(f" ({techs})\n{desc}")

        # Education
        education = doc_data.get("education", [])
        if education:
            doc.add_heading("EDUCATION", level=2)
            for ed in education:
                doc.add_paragraph(f"{ed.get('degree')} in {ed.get('field_of_study', '')} - {ed.get('institution')}")

        doc.save(output_path)
        return output_path
