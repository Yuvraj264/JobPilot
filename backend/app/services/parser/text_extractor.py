import os
from pypdf import PdfReader
from docx import Document


class TextExtractor:
    """
    Text Extraction Service supporting PDF and DOCX documents.
    Detects image-only / scanned PDFs lacking extractable text.
    """

    @staticmethod
    def extract_text(file_path: str, file_type: str) -> str:
        """
        Extracts raw text content from the target document.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Resume file not found at path: {file_path}")

        file_type_upper = file_type.upper()
        if file_type_upper == "PDF":
            return TextExtractor.extract_pdf(file_path)
        elif file_type_upper == "DOCX":
            return TextExtractor.extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported document file type '{file_type}'. Supported types: PDF, DOCX.")

    @staticmethod
    def extract_pdf(file_path: str) -> str:
        """
        Extracts text from PDF using pypdf.
        """
        try:
            reader = PdfReader(file_path)
            extracted_pages = []
            for idx, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    extracted_pages.append(text.strip())

            full_text = "\n\n".join(extracted_pages).strip()
            if not full_text or len(full_text) < 15:
                raise ValueError("No extractable text found in PDF document. Scanned image-only PDFs are not supported without OCR.")

            return full_text
        except ValueError:
            raise
        except Exception as e:
            raise ValueError(f"Failed to read or parse PDF file: {str(e)}")

    @staticmethod
    def extract_docx(file_path: str) -> str:
        """
        Extracts text from DOCX using python-docx.
        """
        try:
            doc = Document(file_path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
            
            # Also extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text and cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)

            full_text = "\n".join(paragraphs).strip()
            if not full_text or len(full_text) < 15:
                raise ValueError("No extractable text found in DOCX document.")

            return full_text
        except ValueError:
            raise
        except Exception as e:
            raise ValueError(f"Failed to read or parse DOCX file: {str(e)}")
